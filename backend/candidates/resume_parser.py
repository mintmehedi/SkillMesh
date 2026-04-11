import re
import unicodedata
from datetime import datetime
from io import BytesIO

from docx import Document

from .ocr_client import extract_text_with_cloud_ocr
from .pdf_extract import extract_pdf_text


def normalize_extracted_resume_text(text: str) -> str:
    """Join hyphenated line wraps and tidy whitespace after PDF/OCR extraction (before heuristics)."""
    if not (text or "").strip():
        return text or ""
    t = unicodedata.normalize("NFKC", text)
    t = re.sub(r"-\s*\r?\n\s*", "", t)
    t = re.sub(r"\u00ad\s*\r?\n?\s*", "", t)
    t = re.sub(r"\r\n|\r", "\n", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    t = re.sub(r"[ \t\f\v]+", " ", t)
    t = re.sub(r" *\n *", "\n", t)
    return t.strip()


# Backward-compatible name (image OCR path).
normalize_post_ocr_text = normalize_extracted_resume_text


COMMON_SKILLS = [
    "python",
    "java",
    "javascript",
    "react",
    "django",
    "sql",
    "postgresql",
    "machine learning",
    "aws",
    "docker",
]

SECTION_HEADER_RE = re.compile(
    r"(work\s+experience|employment\s+history|professional\s+experience|professional\s+background|"
    r"relevant\s+experience|work\s+history|career\s+history|career\s+summary|positions\s+held|"
    r"\bemployment\b|\bexperience\b)",
    re.IGNORECASE,
)
SECTION_STOP_RE = re.compile(
    r"(education|skills|projects|certifications|references|summary|profile)",
    re.IGNORECASE,
)

EDUCATION_HEADER_RE = re.compile(
    r"^(education|academic\s+background|qualifications|academic\s+qualifications)\b",
    re.IGNORECASE,
)
EDUCATION_SECTION_STOP_RE = re.compile(
    r"^(work|employment|experience|professional\s+experience|skills|projects|certifications|references|summary|profile)\b",
    re.IGNORECASE,
)

# Lines inside these blocks must not produce fake "job titles"
SKILL_CERT_BLOCK_START = re.compile(
    r"^(skills?|technical\s+skills|core\s+competencies|certifications?|certificates?|"
    r"licenses?|awards?|languages?|interests|references?)([\s•\-\:]|$)",
    re.IGNORECASE,
)

# Typical job-title tokens (soft requirement when title is short / ambiguous)
_JOB_ROLE_HINT = re.compile(
    r"\b(assistant|engineer|developer|manager|analyst|specialist|consultant|coordinator|"
    r"director|intern|representative|driver|tutor|support|worker|technician|nurse|"
    r"therapist|officer|executive|associate|lead|architect|designer|scientist|"
    r"advisor|administrator|supervisor|facilitator|educator|instructor|marketing|"
    r"photographer|writer|editor|accountant|lawyer|cleaner|cashier|barista|chef|"
    r"volunteer|contractor|freelancer|researcher|plumber|electrician|receptionist)\b",
    re.IGNORECASE,
)


def extract_text_from_upload(uploaded_file):
    name = (uploaded_file.name or "").lower()
    content = uploaded_file.read()
    if hasattr(uploaded_file, "seek"):
        uploaded_file.seek(0)
    if name.endswith(".pdf"):
        return normalize_extracted_resume_text(extract_pdf_text(content))
    if name.endswith(".docx"):
        doc = Document(BytesIO(content))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    parts.append(line)
        return normalize_extracted_resume_text("\n".join(parts))
    if name.endswith(".jpg") or name.endswith(".jpeg") or name.endswith(".png"):
        cloud_text = extract_text_with_cloud_ocr(name, content)
        if cloud_text:
            return normalize_extracted_resume_text(cloud_text)
    return content.decode("utf-8", errors="ignore")


def parse_resume_text(raw_text):
    lowered = raw_text.lower()
    matched_skills = [s for s in COMMON_SKILLS if s in lowered]
    education_level = "bachelor" if "bachelor" in lowered else ("master" if "master" in lowered else "")
    years_experience = 0
    for n in range(15, 0, -1):
        token = f"{n} year"
        if token in lowered:
            years_experience = n
            break
    work_experiences, confidence = _extract_experience_blocks(raw_text)
    studies, edu_conf, edu_section = _extract_education_studies(raw_text)
    return {
        "skills": matched_skills,
        "education_level": education_level,
        "years_experience": years_experience,
        "work_experiences": work_experiences,
        "experience_parse_confidence": confidence,
        "studies": studies,
        "education_parse_confidence": edu_conf,
        "parsed_at": datetime.utcnow().isoformat(),
    }


def _extract_experience_blocks(raw_text):
    lines = [_normalize_extracted_line(line) for line in raw_text.splitlines() if line.strip()]
    lines = _drop_skill_cert_regions(lines)
    exp_lines = _slice_experience_section(lines)
    target_lines = exp_lines if _section_is_useful(exp_lines) else lines

    # Strategy 1: Company|Date anchors (common ATS-friendly pattern)
    s1 = _strategy_company_date_anchor(target_lines)
    # Strategy 2: Date line with nearby title/company context
    s2 = _strategy_date_line_context(target_lines)
    # Strategy 3: Full-document fallback for messy layouts
    s3 = _strategy_date_line_context(lines)

    primary = s1 + s2
    merged = _merge_ranked_experiences(primary if primary else s3)
    experiences = []
    signals = 0
    for idx, row in enumerate(merged[:8]):
        if not _is_valid_title(row["job_title"]):
            continue
        cleaned = {
            "job_title": row["job_title"][:255],
            "company_name": row["company_name"][:255],
            "description": row["description"][:1200],
            "start_date": row["start_date"],
            "end_date": row["end_date"],
            "is_current": row["is_current"],
            "sort_order": idx,
        }
        experiences.append(cleaned)
        signals += int(row.get("_score", 0) * 10)
    confidence = _confidence_score(experiences, signals, bool(exp_lines))
    return experiences, confidence


def _slice_experience_section(lines):
    start_idx = -1
    for i, line in enumerate(lines):
        normalized = _normalize_heading(line)
        if SECTION_HEADER_RE.search(normalized):
            start_idx = i + 1
            break
    if start_idx < 0:
        return []
    out = []
    for line in lines[start_idx:]:
        normalized = _normalize_heading(line)
        if SECTION_STOP_RE.search(normalized):
            break
        out.append(line)
    return out


def _section_is_useful(lines):
    if not lines or len(lines) < 8:
        return False
    date_hits = 0
    for line in lines:
        if _parse_date_range(line) or _parse_company_with_date_line(line):
            date_hits += 1
    return date_hits >= 1


def _normalize_heading(line):
    return re.sub(r"[^a-zA-Z\s]+", "", line or "").strip().lower()


def _drop_skill_cert_regions(lines):
    """Remove lines that belong to Skills / Certifications blocks to reduce false job titles."""
    out = []
    skip = False
    for line in lines:
        stripped = (line or "").strip()
        if not stripped:
            continue
        if SKILL_CERT_BLOCK_START.match(stripped):
            skip = True
            continue
        if skip:
            # End block when we hit a likely new section or a company|date job line
            if SECTION_HEADER_RE.search(_normalize_heading(stripped)):
                skip = False
                continue
            if _parse_company_with_date_line(stripped) or _parse_date_range(stripped):
                skip = False
            else:
                nh = _normalize_heading(stripped)
                if nh in {"education", "experience", "work experience", "employment"}:
                    skip = False
                    continue
                continue
        out.append(line)
    return out


def _normalize_extracted_line(line):
    raw = (line or "").strip()
    if not raw:
        return ""
    # Some PDFs extract as spaced characters:
    # "J u n g s  C l e a n i n g  |  2 0 2 4  M A R - 2 0 2 5  S E P"
    # Use double spaces as word boundaries, then collapse single-char chunks.
    parts = re.split(r"\s{2,}", raw)
    normalized_parts = []
    for part in parts:
        chunk = part.strip()
        if not chunk:
            continue
        tokens = [t for t in chunk.split(" ") if t]
        if len(tokens) >= 2 and all(len(t) == 1 for t in tokens):
            normalized_parts.append("".join(tokens))
        else:
            normalized_parts.append(chunk)
    merged = " ".join(normalized_parts).strip()
    return re.sub(r"\s+", " ", merged)


def _collect_description(lines, start_idx, stop_at_next_job=False):
    parts = []
    for i in range(start_idx, min(start_idx + 14, len(lines))):
        line = lines[i].strip()
        if _parse_date_range(line):
            break
        if stop_at_next_job and _parse_company_with_date_line(line):
            break
        if SECTION_STOP_RE.search(_normalize_heading(line)):
            break
        if len(line) < 2:
            continue
        if _is_valid_title(line) and parts:
            break
        if line.startswith("•") or line.startswith("-"):
            line = line.lstrip("•- ").strip()
        parts.append(line)
    return " ".join(parts)[:1200]


def _strategy_company_date_anchor(lines):
    results = []
    for idx, line in enumerate(lines):
        inline = _parse_company_with_date_line(line)
        if not inline:
            continue
        title_line = _pick_best_title(lines, idx + 1, idx + 2, idx + 3, idx - 1)
        if not title_line:
            title_line = "Unknown Role"
        description_start = idx + 2 if title_line == (lines[idx + 1] if idx + 1 < len(lines) else "") else idx + 1
        description = _collect_description(lines, description_start, stop_at_next_job=True)
        row = {
            "job_title": title_line,
            "company_name": inline["company"],
            "description": description,
            "start_date": inline["start_date"],
            "end_date": inline["end_date"],
            "is_current": inline["is_current"],
            "_score": _entry_score(
                title_line,
                inline["company"],
                inline["start_date"],
                inline["end_date"],
                inline["is_current"],
                description,
                bonus=0.28,
            ),
        }
        results.append(row)
    return results


def _strategy_date_line_context(lines):
    results = []
    for idx, line in enumerate(lines):
        if "|" in line and _parse_company_with_date_line(line):
            continue
        date_info = _parse_date_range(line)
        if not date_info:
            continue

        prev_line = lines[idx - 1] if idx > 0 else ""
        next_line = lines[idx + 1] if idx + 1 < len(lines) else ""
        maybe_after = lines[idx + 2] if idx + 2 < len(lines) else ""

        title, fallback_company = _split_title_company(prev_line)
        if not _is_valid_title(title):
            title = _pick_best_title(lines, idx + 1, idx + 2) or title

        company = ""
        if _looks_like_company(next_line):
            company = next_line
        elif fallback_company and _looks_like_company(fallback_company):
            company = fallback_company
        elif _looks_like_company(maybe_after):
            company = maybe_after

        if not _is_valid_title(title):
            continue

        description_start = idx + 1
        if title and title == next_line:
            description_start = idx + 2
        description = _collect_description(lines, description_start, stop_at_next_job=True)
        row = {
            "job_title": title,
            "company_name": company,
            "description": description,
            "start_date": date_info["start_date"],
            "end_date": date_info["end_date"],
            "is_current": date_info["is_current"],
            "_score": _entry_score(
                title,
                company,
                date_info["start_date"],
                date_info["end_date"],
                date_info["is_current"],
                description,
                bonus=0.06,
            ),
        }
        results.append(row)
    return results


def _pick_best_title(lines, *indices):
    for i in indices:
        if i < 0 or i >= len(lines):
            continue
        candidate = lines[i].strip()
        if _is_valid_title(candidate):
            return candidate
    return ""


def _is_valid_title(text):
    if not text:
        return False
    cleaned = text.strip()
    if len(cleaned) < 3 or len(cleaned) > 120:
        return False
    if cleaned.startswith("•") or cleaned.startswith("-"):
        return False
    if cleaned.endswith("."):
        return False
    if cleaned and cleaned[0].islower():
        return False
    if len(cleaned) <= 8 and cleaned.upper() == cleaned and cleaned.isalpha():
        return False
    # Skill lists: many commas / "and" chains
    if cleaned.count(",") >= 2:
        return False
    if re.search(r"\b(proficient in|familiar with|knowledge of|tools?:)\b", cleaned, flags=re.IGNORECASE):
        return False
    heading = _normalize_heading(cleaned)
    banned = {
        "skills",
        "education",
        "summary",
        "certifications",
        "profile",
        "availability",
        "avabiltiy",
    }
    if heading in banned:
        return False
    if re.search(r"\b(degree|university|college|bachelor|master|phd)\b", cleaned, flags=re.IGNORECASE):
        return False
    if re.search(
        r"\b(police check|national police|ielts|toefl|pmp|aws\s+certified|azure\s+certified|"
        r"comptia|certificate|certification|license number)\b",
        cleaned,
        flags=re.IGNORECASE,
    ):
        return False
    if len(cleaned.split()) > 8:
        return False
    # Short all-caps acronym lines (often certs: RSAT, ACI, etc.) — not job titles
    words = cleaned.split()
    if len(words) == 1 and 2 <= len(cleaned) <= 6 and cleaned.isupper() and cleaned.isalpha():
        return False
    # Prefer titles that look like roles; allow longer lines only if they contain role hints
    if not _JOB_ROLE_HINT.search(cleaned):
        if len(cleaned) < 12:
            return False
        if len(words) <= 2 and not re.search(r"\b(support|care|delivery|domestic|digital)\b", cleaned, re.I):
            return False
    return True


def _entry_score(title, company, start_date, end_date, is_current, description, bonus=0.0):
    score = 0.0
    if _is_valid_title(title):
        score += 0.35
    if company:
        score += 0.22
    if start_date:
        score += 0.22
    if end_date or is_current:
        score += 0.11
    if description and len(description) > 20:
        score += 0.1
    if title.endswith("."):
        score -= 0.12
    if len(title.split()) > 7:
        score -= 0.08
    if not _is_valid_title(title):
        score -= 0.35
    return round(min(1.0, score + bonus), 3)


def _merge_ranked_experiences(rows):
    def key_fn(r):
        return (
            (r.get("job_title") or "").strip().lower(),
            (r.get("start_date") or "").strip(),
        )

    best = {}
    for row in rows:
        k = key_fn(row)
        if not k[0] and not k[1]:
            continue
        cur = best.get(k)
        if (cur is None) or (row.get("_score", 0) > cur.get("_score", 0)):
            best[k] = row

    merged = list(best.values())
    merged.sort(
        key=lambda r: (
            r.get("start_date") or "0000-00-00",
            r.get("_score", 0),
        ),
        reverse=True,
    )
    min_score = 0.34
    return [row for row in merged if row.get("_score", 0) >= min_score]


def _split_title_company(line):
    if " at " in line.lower():
        left, right = re.split(r"\bat\b", line, maxsplit=1, flags=re.IGNORECASE)
        return left.strip(), right.strip()
    if "|" in line:
        left, right = line.split("|", 1)
        return left.strip(), right.strip()
    if "-" in line:
        left, right = line.split("-", 1)
        return left.strip(), right.strip()
    return line.strip(), ""


def _looks_like_company(line):
    lower = line.lower()
    hints = ["ltd", "inc", "pty", "llc", "company", "corp", "technologies", "solutions"]
    bad = [
        "assistant",
        "engineer",
        "developer",
        "driver",
        "support",
        "tutor",
        "student",
    ]
    if any(b in lower for b in bad):
        return False
    return any(h in lower for h in hints) or (len(line.split()) <= 5 and not line.endswith("."))


def _confidence_score(experiences, signals, section_found):
    # No rows => no meaningful score (avoid showing "0% confidence" in the UI).
    if not experiences:
        return None
    score = 0.25
    score += min(0.35, len(experiences) * 0.12)
    score += min(0.15, signals * 0.02)
    if section_found:
        score += 0.08
    return round(min(1.0, score), 2)


def _parse_date_range(line):
    """Parse a date range from one line (or the right-hand side of Company | dates).

    Rules:
    - Year-only tokens use 1 January of that year (start and end).
    - Month + year (no day) use the 1st of that month.
    - Full day + month + year use the exact calendar day.
    - is_current is True only when the end is explicitly Present / Current (word), never as a default.
    """
    split = _split_date_range_tokens(line)
    if not split:
        return None
    start_tok, end_tok, is_present_end = split
    start = _parse_single_date_token(start_tok)
    if not start:
        return None
    if is_present_end:
        return {"start_date": start, "end_date": None, "is_current": True}
    end = _parse_single_date_token(end_tok)
    if not end:
        return {"start_date": start, "end_date": None, "is_current": False}
    return {"start_date": start, "end_date": end, "is_current": False}


def _split_date_range_tokens(text):
    """Return (start_token, end_token_or_None, is_present_end) or None."""
    text = (text or "").strip()
    if not text:
        return None
    # Explicit Present / Current after a dash (handles "Jan 2020 - Present", "2022-Present")
    m = re.search(r"(?:\s*[-–]\s*)(present|current)\s*$", text, re.IGNORECASE)
    if m:
        before = text[: m.start()].strip()
        if before:
            return (before, None, True)
    # "Jan 2020 to Mar 2022" / "2020 to Present"
    if re.search(r"\bto\b", text, re.IGNORECASE):
        m2 = re.search(r"\bto\s+(present|current)\s*$", text, re.IGNORECASE)
        if m2:
            before = text[: m2.start()].strip()
            if before:
                return (before, None, True)
        parts = re.split(r"\s+to\s+", text, maxsplit=1, flags=re.IGNORECASE)
        if len(parts) == 2:
            a, b = parts[0].strip(), parts[1].strip()
            if a and b:
                if re.match(r"^(present|current)$", b, re.IGNORECASE):
                    return (a, None, True)
                return (a, b, False)
    # Two years only: 2020-2022 (do not split single hyphens inside words)
    ym = re.match(r"^(\d{4})\s*[-–]\s*(\d{4})$", text)
    if ym:
        return (ym.group(1), ym.group(2), False)
    # Spaced dash avoids breaking hyphenated words (e.g. Non-current … 2020 - 2022)
    if re.search(r"\s[-–]\s", text):
        parts = re.split(r"\s+[-–]\s+", text, maxsplit=1)
        if len(parts) == 2:
            a, b = parts[0].strip(), parts[1].strip()
            if a and b:
                if re.match(r"^(present|current)$", b, re.IGNORECASE):
                    return (a, None, True)
                return (a, b, False)
    parts = re.split(r"\s*[-–]\s*", text, maxsplit=1)
    if len(parts) != 2:
        return None
    a, b = parts[0].strip(), parts[1].strip()
    if not a or not b:
        return None
    if re.match(r"^(present|current)$", b, re.IGNORECASE):
        return (a, None, True)
    return (a, b, False)


def _parse_single_date_token(token):
    """Normalize one date token to YYYY-MM-DD, or ''."""
    if not token:
        return ""
    t = re.sub(r"\s+", " ", token.strip())
    if not t:
        return ""

    # ISO YYYY-MM-DD
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})$", t)
    if m:
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if 1 <= mo <= 12 and 1 <= d <= 31:
            return f"{m.group(1)}-{m.group(2)}-{m.group(3)}"

    # Year only → 1 Jan
    if re.match(r"^\d{4}$", t):
        return f"{t}-01-01"

    # YYYY MM (e.g. 2024 MAR)
    m = re.match(r"^(\d{4})\s+([A-Za-z]{3,9})$", t)
    if m:
        mo = _month_to_number(m.group(2))
        return f"{m.group(1)}-{mo}-01"

    # Month YYYY (e.g. Mar 2024, September 2024)
    m = re.match(r"^([A-Za-z]{3,9})\s+(\d{4})$", t)
    if m:
        mo = _month_to_number(m.group(1))
        return f"{m.group(2)}-{mo}-01"

    # DD Month YYYY
    m = re.match(r"^(\d{1,2})\s+([A-Za-z]{3,9})\s+(\d{4})$", t)
    if m:
        d, mon, y = int(m.group(1)), m.group(2), m.group(3)
        mo = _month_to_number(mon)
        if _is_valid_calendar_day(y, mo, d):
            return f"{y}-{mo}-{int(d):02d}"

    # Month DD, YYYY or Month DD YYYY
    m = re.match(r"^([A-Za-z]{3,9})\s+(\d{1,2}),?\s+(\d{4})$", t)
    if m:
        mon, d, y = m.group(1), int(m.group(2)), m.group(3)
        mo = _month_to_number(mon)
        if _is_valid_calendar_day(y, mo, d):
            return f"{y}-{mo}-{int(d):02d}"

    # MM/YYYY or M/YYYY
    m = re.match(r"^(\d{1,2})/(\d{4})$", t)
    if m:
        mo, y = int(m.group(1)), m.group(2)
        if 1 <= mo <= 12:
            return f"{y}-{mo:02d}-01"

    # DD/MM/YYYY (common outside US)
    m = re.match(r"^(\d{1,2})/(\d{1,2})/(\d{4})$", t)
    if m:
        a, b, y = int(m.group(1)), int(m.group(2)), m.group(3)
        if a > 12:
            d, mo = a, b
        elif b > 12:
            mo, d = a, b
        else:
            d, mo = a, b
        if 1 <= mo <= 12 and _is_valid_calendar_day(y, f"{mo:02d}", d):
            return f"{y}-{mo:02d}-{int(d):02d}"

    return ""


def _is_valid_calendar_day(year_str, month_str, day_int):
    try:
        y = int(year_str)
        mo = int(month_str) if isinstance(month_str, str) and month_str.isdigit() else int(month_str)
        datetime(y, mo, int(day_int))
        return True
    except (ValueError, TypeError):
        return False


def _month_to_number(raw):
    m = raw.strip().lower()[:3]
    table = {
        "jan": "01",
        "feb": "02",
        "mar": "03",
        "apr": "04",
        "may": "05",
        "jun": "06",
        "jul": "07",
        "aug": "08",
        "sep": "09",
        "oct": "10",
        "nov": "11",
        "dec": "12",
    }
    return table.get(m, "01")


def _parse_company_with_date_line(line):
    # Examples: "Jungs Cleaning Ryde | 2024 MAR - 2025 SEP"
    # ATS columns: "Acme Pty Ltd    Jan 2020 - Dec 2022"
    # Em/en dash: "TechCorp — Mar 2019 - Present"
    line = (line or "").strip()
    if not line:
        return None

    for sep in ("\u2014", "\u2013"):  # em dash, en dash
        if sep in line:
            left, right = line.split(sep, 1)
            left, right = left.strip(), right.strip()
            date_info = _parse_date_range(right)
            if date_info and left and not _parse_date_range(left):
                return {
                    "company": left,
                    "start_date": date_info["start_date"],
                    "end_date": date_info["end_date"],
                    "is_current": date_info["is_current"],
                }

    chunks = [c.strip() for c in re.split(r"\s{2,}", line) if c.strip()]
    if len(chunks) == 2:
        left, right = chunks[0], chunks[1]
        if not _parse_date_range(left):
            date_info = _parse_date_range(right)
            if date_info and left:
                return {
                    "company": left,
                    "start_date": date_info["start_date"],
                    "end_date": date_info["end_date"],
                    "is_current": date_info["is_current"],
                }

    if "|" not in line:
        return None
    left, right = line.split("|", 1)
    date_info = _parse_date_range(right.strip())
    if not date_info:
        return None
    company = left.strip()
    if not company:
        return None
    return {
        "company": company,
        "start_date": date_info["start_date"],
        "end_date": date_info["end_date"],
        "is_current": date_info["is_current"],
    }


def _extract_education_studies(raw_text):
    """Pull education/study blocks from an Education section (heuristic)."""
    lines = [_normalize_extracted_line(line) for line in raw_text.splitlines() if line.strip()]
    start_idx = -1
    for i, line in enumerate(lines):
        if EDUCATION_HEADER_RE.search(_normalize_heading(line)):
            start_idx = i + 1
            break
    if start_idx < 0:
        return [], None, False
    section_lines = []
    for line in lines[start_idx:]:
        if EDUCATION_SECTION_STOP_RE.search(_normalize_heading(line)):
            break
        section_lines.append(line)
    if not section_lines:
        return [], None, True
    blob = "\n".join(section_lines)
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n+", blob) if p.strip()]
    studies = []
    for para in paragraphs[:8]:
        row = _parse_education_paragraph(para)
        if row:
            row["sort_order"] = len(studies)
            studies.append(row)
    conf = None
    if studies:
        conf = round(min(1.0, 0.22 + min(0.5, len(studies) * 0.12) + 0.08), 2)
    return studies, conf, True


def _parse_education_paragraph(para):
    one_line = re.sub(r"\s+", " ", para).strip()
    if len(one_line) < 6:
        return None
    institution = ""
    m = re.search(
        r"([A-Z0-9][A-Za-z0-9\s,'&\-.]+?(?:University|College|Institute|School|Academy))\b",
        para,
        re.I,
    )
    if m:
        institution = re.sub(r"\s+", " ", m.group(1).strip())[:255]
    degree = ""
    dm = re.search(
        r"(Bachelor(?:'s)?(?:\s+of)?[^,\n|]*|B\.?S\.?c?\.?|B\.?A\.?|Master(?:'s)?(?:\s+of)?[^,\n|]*|M\.?S\.?c?\.?|M\.?A\.?|MBA|Ph\.?D\.?|Doctorate|Doctor\s+of[^,\n|]*|Diploma|Associate(?:'s)?[^,\n|]*|Graduate\s+Certificate[^,\n|]*)",
        para,
        re.I,
    )
    if dm:
        degree = re.sub(r"\s+", " ", dm.group(0).strip())[:255]
    field_of_study = ""
    in_m = re.search(
        r"\b(?:in|of)\s+([A-Za-z0-9][A-Za-z0-9\s\-&/.]+?)(?=\s*[|,]|\s+at\s|\s+—|\s+–|\s+\d{4}|\s*$)",
        para,
        re.I,
    )
    if in_m:
        field_of_study = in_m.group(1).strip()[:255]
    major = field_of_study
    mj = re.search(r"Majors?[:\s]+([^,|\n]+)", para, re.I)
    if mj:
        major = mj.group(1).strip()[:255]
    dr = _parse_date_range(one_line)
    start_d = dr["start_date"] if dr else ""
    end_d = "" if (dr and dr["is_current"]) else (dr["end_date"] if dr else "")
    is_cur = bool(dr["is_current"]) if dr else False
    if not institution and not degree and not field_of_study and not major:
        if len(one_line) < 20:
            return None
        institution = one_line[:255]
    desc = para.strip()[:1200] if len(para) > 200 else ""
    return {
        "institution": institution,
        "degree": degree,
        "field_of_study": field_of_study,
        "major": major,
        "description": desc,
        "start_date": start_d or "",
        "end_date": end_d or "",
        "is_current": is_cur,
        "sort_order": 0,
    }
